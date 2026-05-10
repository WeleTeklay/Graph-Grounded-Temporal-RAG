#src/document_loader.py
"""
Multi-format document loader supporting PDF, DOCX, TXT, MD, HTML.
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

import re
from typing import Optional, Dict, Any, List
from dataclasses import dataclass

from src.logger import get_logger
from src.utils import handle_errors

logger = get_logger(__name__)


@dataclass
class LoadedDocument:
    """Structured loaded document."""
    doc_id: str
    file_path: Path
    file_type: str
    text: str
    metadata: Dict[str, Any]
    page_count: int
    char_count: int
    word_count: int


class MultiFormatLoader:
    """
    Universal document loader supporting multiple formats.
    """
    
    SUPPORTED_FORMATS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
        '.txt': 'text',
        '.md': 'markdown',
        '.html': 'html',
        '.htm': 'html',
        '.rtf': 'rtf',
        '.odt': 'odt',
    }
    
    def __init__(self):
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check which format libraries are available."""
        self.has_pdf = self._check_import('fitz', 'PyMuPDF')
        self.has_docx = self._check_import('docx', 'python-docx')
        self.has_antiword = self._check_import('antiword', 'antiword')
        self.has_odt = self._check_import('odf', 'odfpy')
        
        logger.info(f"Format support: PDF={self.has_pdf}, DOCX={self.has_docx}")
    
    def _check_import(self, module: str, package: str) -> bool:
        try:
            __import__(module)
            return True
        except ImportError:
            logger.warning(f"{package} not installed. {module} support disabled.")
            return False
    
    def get_file_type(self, file_path: Path) -> str:
        """Detect file type from extension."""
        ext = file_path.suffix.lower()
        return self.SUPPORTED_FORMATS.get(ext, 'unknown')
    
    @handle_errors(default_return=None)
    def load_pdf(self, file_path: Path) -> Optional[LoadedDocument]:
        """Load PDF document."""
        if not self.has_pdf:
            logger.error("PyMuPDF not installed. Run: pip install pymupdf")
            return None
        
        import fitz
        
        doc = fitz.open(file_path)
        text_parts = []
        metadata = dict(doc.metadata)
        
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"[PAGE {page_num}]\n{page_text}")
        
        doc.close()
        
        full_text = "\n\n".join(text_parts)
        
        return LoadedDocument(
            doc_id=file_path.stem,
            file_path=file_path,
            file_type='pdf',
            text=full_text,
            metadata=metadata,
            page_count=len(doc),
            char_count=len(full_text),
            word_count=len(full_text.split())
        )
    
    @handle_errors(default_return=None)
    def load_docx(self, file_path: Path) -> Optional[LoadedDocument]:
        """Load DOCX document."""
        if not self.has_docx:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return None
        
        from docx import Document
        
        doc = Document(file_path)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text)
                if row_text:
                    text_parts.append(row_text)
        
        full_text = "\n\n".join(text_parts)
        
        return LoadedDocument(
            doc_id=file_path.stem,
            file_path=file_path,
            file_type='docx',
            text=full_text,
            metadata={'core_properties': {}},
            page_count=len(text_parts) // 40 + 1,
            char_count=len(full_text),
            word_count=len(full_text.split())
        )
    
    @handle_errors(default_return=None)
    def load_text(self, file_path: Path) -> Optional[LoadedDocument]:
        """Load plain text document with encoding detection."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                break
            except UnicodeDecodeError:
                continue
        else:
            logger.error(f"Could not decode {file_path}")
            return None
        
        return LoadedDocument(
            doc_id=file_path.stem,
            file_path=file_path,
            file_type='text',
            text=text,
            metadata={'encoding': encoding},
            page_count=1,
            char_count=len(text),
            word_count=len(text.split())
        )
    
    @handle_errors(default_return=None)
    def load_markdown(self, file_path: Path) -> Optional[LoadedDocument]:
        """Load Markdown document."""
        import markdown
        from bs4 import BeautifulSoup
        
        with open(file_path, 'r', encoding='utf-8') as f:
            md_text = f.read()
        
        # Convert to HTML then extract text
        html = markdown.markdown(md_text)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text(separator='\n')
        
        return LoadedDocument(
            doc_id=file_path.stem,
            file_path=file_path,
            file_type='markdown',
            text=text,
            metadata={'format': 'markdown'},
            page_count=1,
            char_count=len(text),
            word_count=len(text.split())
        )
    
    @handle_errors(default_return=None)
    def load_html(self, file_path: Path) -> Optional[LoadedDocument]:
        """Load HTML document."""
        from bs4 import BeautifulSoup
        
        with open(file_path, 'r', encoding='utf-8') as f:
            html = f.read()
        
        soup = BeautifulSoup(html, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text(separator='\n')
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        text = '\n'.join(lines)
        
        # Extract metadata
        title = soup.title.string if soup.title else file_path.stem
        
        return LoadedDocument(
            doc_id=file_path.stem,
            file_path=file_path,
            file_type='html',
            text=text,
            metadata={'title': title},
            page_count=1,
            char_count=len(text),
            word_count=len(text.split())
        )
    
    def load(self, file_path: Path) -> Optional[LoadedDocument]:
        """
        Load any supported document format.
        
        Args:
            file_path: Path to document file
        
        Returns:
            LoadedDocument or None if unsupported/failed
        """
        file_type = self.get_file_type(file_path)
        
        loaders = {
            'pdf': self.load_pdf,
            'docx': self.load_docx,
            'doc': self.load_docx,  # Same as docx
            'text': self.load_text,
            'markdown': self.load_markdown,
            'html': self.load_html,
        }
        
        loader = loaders.get(file_type)
        if not loader:
            logger.error(f"Unsupported format: {file_type}")
            return None
        
        logger.info(f"Loading {file_type.upper()}: {file_path.name}")
        return loader(file_path)
    
    def load_all(self, directory: Path, pattern: str = "*") -> List[LoadedDocument]:
        """
        Load all supported documents in a directory.
        
        Args:
            directory: Directory containing documents
            pattern: Glob pattern (e.g., "*.pdf", "*.*")
        
        Returns:
            List of successfully loaded documents
        """
        documents = []
        
        for ext in self.SUPPORTED_FORMATS.keys():
            for file_path in directory.glob(f"*{ext}"):
                if pattern != "*" and not file_path.match(pattern):
                    continue
                
                doc = self.load(file_path)
                if doc:
                    documents.append(doc)
        
        logger.info(f"Loaded {len(documents)} documents from {directory}")
        return documents