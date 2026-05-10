"""
Phase 1: Production-Grade Universal Document Parser
Supports PDF, DOCX, TXT, MD, HTML with metadata preservation and OCR fallback.
ADDED: Language detection and English-only extraction for bilingual PDFs.
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

import os
import re
import json
import hashlib
from datetime import datetime
from typing import Dict, Optional, Tuple, List, Union
from dataclasses import dataclass, asdict

import pandas as pd
import fitz  # PyMuPDF

from src.logger import get_logger
from src.config import config
from src.utils import handle_errors, safe_file_read, safe_json_dump, Timer

logger = get_logger(__name__)

# ============================================================================
# Optional Format Support (Graceful Degradation)
# ============================================================================
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.debug("python-docx not installed. DOCX support disabled.")

try:
    import markdown
    from bs4 import BeautifulSoup
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False
    logger.debug("markdown/beautifulsoup4 not installed. MD/HTML support disabled.")


# ============================================================================
# Supported Formats
# ============================================================================
SUPPORTED_EXTENSIONS = {
    '.pdf': 'pdf',
    '.docx': 'docx',
    '.doc': 'docx',
    '.txt': 'text',
    '.md': 'markdown',
    '.html': 'html',
    '.htm': 'html',
}

# Amharic Unicode range: U+1200 to U+137F
# Also includes Ethiopic supplement U+1380 to U+139F, U+2D80 to U+2DDF
AMHARIC_PATTERN = re.compile(r'[\u1200-\u137F\u1380-\u139F\u2D80-\u2DDF]+')


@dataclass
class ParsedDocument:
    """Structured representation of a parsed document."""
    doc_id: str
    doc_title: str
    effective_date: str
    supersedes_doc_id: Optional[str]
    source_path: str
    output_path: str
    file_type: str
    page_count: int
    char_count: int
    word_count: int
    file_hash: str
    parsed_at: str
    status: str
    error_message: Optional[str] = None


class UniversalDocumentParser:
    """
    Production-grade universal document parser supporting multiple formats.
    Maintains 100% backward compatibility with PDF-only mode.
    """
    
    def __init__(self):
        self.manifest_path = config.paths.project_root / "document_manifest.csv"
        self.raw_docs_dir = config.paths.raw_pdfs_dir
        self.processed_dir = config.paths.processed_texts_dir
        self.processed_dir.mkdir(parents=True, exist_ok=True)
        
        # Format support flags
        self.has_docx = HAS_DOCX
        self.has_markdown = HAS_MARKDOWN
        
        logger.info(f"Format support: PDF=YES, DOCX={self.has_docx}, Markdown/HTML={self.has_markdown}")
    
    def _find_document_file(self, doc_id: str) -> Optional[Path]:
        """Find document file with any supported extension."""
        priority_order = ['.pdf', '.docx', '.doc', '.txt', '.md', '.html', '.htm']
        
        for ext in priority_order:
            candidate = self.raw_docs_dir / f"{doc_id}{ext}"
            if candidate.exists():
                return candidate
        
        return None
    
    def _get_file_type(self, file_path: Path) -> str:
        """Determine file type from extension."""
        ext = file_path.suffix.lower()
        return SUPPORTED_EXTENSIONS.get(ext, 'unknown')
    
    def _filter_english_only(self, text: str) -> str:
        """
        Filter out non-English text (specifically Amharic characters).
        Keeps only ASCII/English text.
        """
        if not text:
            return ""
        
        # Method 1: Remove Amharic characters
        text = AMHARIC_PATTERN.sub('', text)
        
        # Method 2: Keep only ASCII printable characters + common punctuation
        # This removes other non-English scripts too
        ascii_text = re.sub(r'[^\x00-\x7F]+', ' ', text)
        
        # Method 3: Additional cleaning - remove multiple spaces
        ascii_text = re.sub(r'\s+', ' ', ascii_text)
        
        # Remove lines that are too short (likely remnants of removed text)
        lines = ascii_text.split('\n')
        filtered_lines = []
        for line in lines:
            # Keep lines with at least 20 characters and some English words
            if len(line.strip()) > 20 and re.search(r'[a-zA-Z]{3,}', line):
                filtered_lines.append(line)
        
        return '\n'.join(filtered_lines)
    
    def _detect_and_extract_english_column(self, page) -> str:
        """
        Detect two-column layout and extract only the English column.
        Assumes English is on the right side.
        """
        # Get page text with position information
        page_dict = page.get_text("dict")
        blocks = page_dict.get("blocks", [])
        
        if not blocks:
            return page.get_text()
        
        # Collect x-coordinates of text blocks
        x_coords = []
        for block in blocks:
            if "bbox" in block:
                x_coords.append(block["bbox"][0])  # x0 coordinate
        
        if len(x_coords) < 5:
            # Not enough data for column detection, return full text
            return page.get_text()
        
        # Find clusters of x-coordinates
        from collections import Counter
        x_rounded = [round(x / 100) * 100 for x in x_coords]
        x_counts = Counter(x_rounded)
        
        # Get the two most common x-positions (left and right columns)
        most_common = x_counts.most_common(2)
        
        if len(most_common) >= 2:
            # Two columns detected
            left_col_x = most_common[0][0]
            right_col_x = most_common[1][0]
            
            # Ensure right column is actually to the right
            if left_col_x > right_col_x:
                left_col_x, right_col_x = right_col_x, left_col_x
            
            # Extract only right column text (assumed English)
            right_column_text = []
            
            for block in blocks:
                if "bbox" in block and "lines" in block:
                    x0 = block["bbox"][0]
                    # Check if block is in right column
                    if abs(x0 - right_col_x) < 100:
                        for line in block["lines"]:
                            for span in line.get("spans", []):
                                text = span["text"].strip()
                                if text:
                                    right_column_text.append(text)
            
            if right_column_text:
                return " ".join(right_column_text)
        
        # Fallback: return full text and let language filter handle it
        return page.get_text()
    
    @handle_errors(default_return=(None, {}))
    def _load_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Load and extract text from PDF with bilingual support."""
        doc = fitz.open(file_path)
        
        metadata = {
            "page_count": len(doc),
            "pdf_title": doc.metadata.get("title", ""),
            "pdf_author": doc.metadata.get("author", ""),
            "pdf_subject": doc.metadata.get("subject", ""),
            "pdf_creator": doc.metadata.get("creator", ""),
            "pdf_producer": doc.metadata.get("producer", ""),
            "pdf_creation_date": doc.metadata.get("creationDate", ""),
        }
        
        full_text = []
        for page_num, page in enumerate(doc, 1):
            # Try column detection first
            page_text = self._detect_and_extract_english_column(page)
            
            # If column detection gave little text, try normal extraction
            if len(page_text.strip()) < 100:
                page_text = page.get_text()
            
            # Filter out non-English (Amharic) characters
            page_text = self._filter_english_only(page_text)
            
            if page_text.strip():
                full_text.append(f"[PAGE {page_num}]\n{page_text}")
        
        doc.close()
        
        combined_text = "\n\n".join(full_text)
        
        # Check if PDF is scanned
        if len(combined_text.strip()) < 100:
            logger.warning(f"Possible scanned PDF detected: {file_path.name}")
            combined_text = self._attempt_ocr_fallback(file_path)
            # Also filter OCR results
            combined_text = self._filter_english_only(combined_text)
        
        return combined_text, metadata
    
    @handle_errors(default_return=(None, {}))
    def _load_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """Load and extract text from DOCX with language filtering."""
        if not self.has_docx:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = Document(file_path)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                filtered = self._filter_english_only(para.text)
                if filtered.strip():
                    text_parts.append(filtered)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                filtered = self._filter_english_only(row_text)
                if filtered.strip():
                    text_parts.append(filtered)
        
        full_text = "\n\n".join(text_parts)
        
        metadata = {
            "page_count": len(doc.paragraphs) // 40 + 1,
            "doc_title": file_path.stem,
            "doc_author": "",
        }
        
        return full_text, metadata
    
    @handle_errors(default_return=(None, {}))
    def _load_text(self, file_path: Path) -> Tuple[str, Dict]:
        """Load plain text file with encoding detection and language filtering."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError(f"Could not decode {file_path}")
        
        # Filter out non-English
        text = self._filter_english_only(text)
        
        metadata = {
            "page_count": 1,
            "doc_title": file_path.stem,
            "encoding": encoding,
        }
        
        return text, metadata
    
    @handle_errors(default_return=(None, {}))
    def _load_markdown(self, file_path: Path) -> Tuple[str, Dict]:
        """Load and extract text from Markdown with language filtering."""
        if not self.has_markdown:
            raise ImportError("markdown/beautifulsoup4 not installed. Run: pip install markdown beautifulsoup4")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            md_text = f.read()
        
        html = markdown.markdown(md_text)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text(separator='\n')
        
        # Filter out non-English
        text = self._filter_english_only(text)
        
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        text = '\n'.join(lines)
        
        metadata = {
            "page_count": 1,
            "doc_title": file_path.stem,
            "format": "markdown",
        }
        
        return text, metadata
    
    @handle_errors(default_return=(None, {}))
    def _load_html(self, file_path: Path) -> Tuple[str, Dict]:
        """Load and extract text from HTML with language filtering."""
        if not self.has_markdown:
            raise ImportError("beautifulsoup4 not installed. Run: pip install beautifulsoup4")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            html = f.read()
        
        soup = BeautifulSoup(html, 'html.parser')
        
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text(separator='\n')
        
        # Filter out non-English
        text = self._filter_english_only(text)
        
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        text = '\n'.join(lines)
        
        title = soup.title.string if soup.title else file_path.stem
        
        metadata = {
            "page_count": 1,
            "doc_title": title,
            "format": "html",
        }
        
        return text, metadata
    
    def extract_text_with_metadata(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text and metadata from any supported document format."""
        file_type = self._get_file_type(file_path)
        
        loaders = {
            'pdf': self._load_pdf,
            'docx': self._load_docx,
            'text': self._load_text,
            'markdown': self._load_markdown,
            'html': self._load_html,
        }
        
        loader = loaders.get(file_type)
        if not loader:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        logger.debug(f"Loading {file_type.upper()}: {file_path.name}")
        return loader(file_path)
    
    def _attempt_ocr_fallback(self, pdf_path: Path) -> str:
        """Attempt OCR for scanned PDFs with language filtering."""
        try:
            from pdf2image import convert_from_path
            import pytesseract
            
            logger.info(f"Attempting OCR on {pdf_path.name}...")
            images = convert_from_path(pdf_path, dpi=150)
            
            ocr_text = []
            for i, image in enumerate(images, 1):
                text = pytesseract.image_to_string(image)
                text = self._filter_english_only(text)
                ocr_text.append(f"[PAGE {i} - OCR]\n{text}")
            
            return "\n\n".join(ocr_text)
            
        except ImportError:
            logger.warning("OCR dependencies not installed. Install: pdf2image, pytesseract")
            return "[SCANNED PDF - OCR NOT AVAILABLE]"
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return "[SCANNED PDF - OCR FAILED]"
    
    def _compute_file_hash(self, file_path: Path) -> str:
        """Compute SHA-256 hash of file."""
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        return sha256.hexdigest()
    
    def _count_words(self, text: str) -> int:
        """Count words in extracted text."""
        return len(text.split())
    
    def _create_metadata_header(self, doc_id: str, effective_date: str, 
                                 supersedes: Optional[str], file_type: str,
                                 metadata: Dict) -> str:
        """Create structured metadata header for processed text file."""
        header = f"""[DOCUMENT METADATA]
DOC_ID: {doc_id}
EFFECTIVE_DATE: {effective_date}
SUPERSEDES: {supersedes if supersedes and pd.notna(supersedes) else 'None'}
FILE_TYPE: {file_type}
PAGE_COUNT: {metadata.get('page_count', 0)}
DOC_TITLE: {metadata.get('doc_title', metadata.get('pdf_title', 'N/A'))}
DOC_AUTHOR: {metadata.get('doc_author', metadata.get('pdf_author', 'N/A'))}
PROCESSED_AT: {datetime.now().isoformat()}
[END METADATA]

"""
        return header
    
    @handle_errors(default_return=None)
    def parse_single_document(self, doc_id: str, doc_title: str, 
                               effective_date: str, supersedes: Optional[str]) -> Optional[ParsedDocument]:
        """Parse a single document (any supported format)."""
        file_path = self._find_document_file(doc_id)
        
        if not file_path:
            logger.error(f"Document not found: {doc_id}.* in {self.raw_docs_dir}")
            return None
        
        file_type = self._get_file_type(file_path)
        logger.info(f"Parsing: {doc_id} - {doc_title} ({file_type.upper()})")
        
        with Timer(f"Parse {doc_id}"):
            extracted_text, metadata = self.extract_text_with_metadata(file_path)
            
            if not extracted_text:
                logger.error(f"No text extracted from {doc_id}")
                return None
            
            metadata['doc_title'] = doc_title
            
            char_count = len(extracted_text)
            word_count = self._count_words(extracted_text)
            file_hash = self._compute_file_hash(file_path)
            
            header = self._create_metadata_header(
                doc_id, effective_date, supersedes, file_type, metadata
            )
            
            output_content = header + "[TEXT CONTENT]\n\n" + extracted_text
            output_path = self.processed_dir / f"{doc_id}.txt"
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(output_content)
            
            logger.info(f"Parsed {doc_id}: {metadata.get('page_count', 1)} pages, "
                       f"{char_count:,} chars, {word_count:,} words")
            
            return ParsedDocument(
                doc_id=doc_id,
                doc_title=doc_title,
                effective_date=effective_date,
                supersedes_doc_id=supersedes if pd.notna(supersedes) else None,
                source_path=str(file_path),
                output_path=str(output_path),
                file_type=file_type,
                page_count=metadata.get('page_count', 1),
                char_count=char_count,
                word_count=word_count,
                file_hash=file_hash,
                parsed_at=datetime.now().isoformat(),
                status="success"
            )
    
    def parse_all_documents(self) -> List[ParsedDocument]:
        """Parse all documents defined in manifest."""
        if not self.manifest_path.exists():
            logger.error(f"Manifest not found: {self.manifest_path}")
            return []
        
        manifest = pd.read_csv(self.manifest_path)
        logger.info(f"Found {len(manifest)} documents in manifest")
        
        parsed_docs = []
        failed_docs = []
        
        for _, row in manifest.iterrows():
            doc_id = row['doc_id']
            doc_title = row['doc_title']
            effective_date = row['effective_date']
            supersedes = row.get('supersedes_doc_id', None)
            
            result = self.parse_single_document(doc_id, doc_title, effective_date, supersedes)
            
            if result:
                parsed_docs.append(result)
            else:
                failed_docs.append({
                    'doc_id': doc_id,
                    'doc_title': doc_title,
                    'error': 'Parsing failed'
                })
        
        report = {
            'parsed_at': datetime.now().isoformat(),
            'total_documents': len(manifest),
            'successful': len(parsed_docs),
            'failed': len(failed_docs),
            'parsed_documents': [asdict(d) for d in parsed_docs],
            'failed_documents': failed_docs
        }
        
        report_path = self.processed_dir / "parsing_report.json"
        safe_json_dump(report, report_path)
        
        logger.info(f"Parsing complete: {len(parsed_docs)}/{len(manifest)} successful")
        
        if failed_docs:
            logger.warning(f"Failed documents: {[d['doc_id'] for d in failed_docs]}")
        
        return parsed_docs


# ============================================================================
# Backward Compatibility Alias
# ============================================================================
PDFParser = UniversalDocumentParser


def main():
    """Main entry point for universal document parsing phase."""
    logger.info("=" * 60)
    logger.info("Phase 1: Production Universal Document Parser")
    logger.info("=" * 60)
    
    parser = UniversalDocumentParser()
    parsed_docs = parser.parse_all_documents()
    
    if parsed_docs:
        logger.info(f"✅ Successfully parsed {len(parsed_docs)} documents")
        logger.info(f"📁 Output directory: {config.paths.processed_texts_dir}")
        logger.info(f"📊 Report: {config.paths.processed_texts_dir}/parsing_report.json")
    else:
        logger.error("❌ No documents were parsed successfully")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())